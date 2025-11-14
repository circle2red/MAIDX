# TODO: make this a class and make extraction_worker use this class
# use class stores config
"""
File parsers for various document formats
"""
from typing import Optional
import os


def parse_txt(file_path: str) -> str:
    """
    Parse a text file

    Args:
        file_path: Path to the text file

    Returns:
        Text content
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        # Try with different encoding
        with open(file_path, 'r', encoding='latin-1') as f:
            return f.read()


def parse_pdf(file_path: str) -> str:
    """
    Parse a PDF file

    Args:
        file_path: Path to the PDF file

    Returns:
        Extracted text content
    """
    try:
        import pypdf

        text = []
        with open(file_path, 'rb') as f:
            pdf_reader = pypdf.PdfReader(f)
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)

        return '\n'.join(text)

    except ImportError:
        return "[Error: PyPDF not installed. Install with: pip install PyPDF]"
    except Exception as e:
        return f"[Error parsing PDF: {str(e)}]"


def parse_docx(file_path: str) -> str:
    """
    Parse a DOCX file

    Args:
        file_path: Path to the DOCX file

    Returns:
        Extracted text content
    """
    try:
        from docx import Document

        doc = Document(file_path)
        text = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text.append(row_text)

        return '\n'.join(text)

    except ImportError:
        return "[Error: python-docx not installed. Install with: pip install python-docx]"
    except Exception as e:
        return f"[Error parsing DOCX: {str(e)}]"


def parse_image(file_path: str) -> str:
    """
    Parse an image file using OCR

    Args:
        file_path: Path to the image file

    Returns:
        Extracted text content
    """
    try:
        from PIL import Image
        import pytesseract

        image = Image.open(file_path)
        text = pytesseract.image_to_string(image)

        return text.strip()

    except ImportError as e:
        if 'PIL' in str(e) or 'Image' in str(e):
            return "[Error: Pillow not installed. Install with: pip install Pillow]"
        elif 'pytesseract' in str(e):
            return "[Error: pytesseract not installed. Install with: pip install pytesseract]"
        else:
            return f"[Error: Missing dependency - {str(e)}]"
    except Exception as e:
        return f"[Error parsing image: {str(e)}]"


def parse_file(file_path: str) -> Optional[str]:
    """
    Parse a file based on its extension

    Args:
        file_path: Path to the file

    Returns:
        Extracted text content or None if unsupported
    """
    ext = os.path.splitext(file_path)[1].lower()

    parsers = {
        '.txt': parse_txt,
        '.pdf': parse_pdf,
        '.docx': parse_docx,
        '.jpg': parse_image,
        '.jpeg': parse_image,
        '.png': parse_image,
    }

    parser = parsers.get(ext)
    if parser:
        return parser(file_path)
    else:
        return None


def test_parsers():
    """Test the file parsers"""
    import tempfile

    # Test TXT parser
    with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False) as f:
        f.write("Hello, world!\nThis is a test.")
        txt_path = f.name

    print("TXT content:", parse_txt(txt_path))
    os.unlink(txt_path)

    print("\nOther parsers require actual files to test.")


if __name__ == "__main__":
    test_parsers()
